# -*- coding: utf-8 -*-
"""Generate 清竞数智 platform intro Word doc."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)

    title = doc.add_heading("清竞数智 · 高质量数据集建设平台", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in title.runs:
        r.font.name = "黑体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        r.font.size = Pt(22)

    sub = doc.add_paragraph("功能介绍（自主研发）")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.font.size = Pt(12)
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("一、平台概述").bold = True
    body = (
        "清竞数智高质量数据集建设平台由我司自主研发，面向大语言模型全生命周期，提供从原始资料到可用模型的"
        "一体化数据建设与训练编排能力。平台产品形态、工作流编排、权限与审计、对外接口及交付运维体系均为"
        "清竞数智自主设计与研发；在统一工作台内完成数据治理、多阶段训练任务配置与执行、"
        "以及模型与数据效果的闭环评测，支撑企业级高质量数据集资产沉淀与模型能力持续迭代。\n\n"
        "在工程实现上，平台深度融合业界成熟的开源技术栈：在数据侧集成 Easy Dataset 的文档解析、智能分块、"
        "问答与对话合成、人机协同质检等能力；在训练侧对接 LLaMA-Factory 的统一训练框架，覆盖预训练、"
        "监督微调及强化学习与偏好对齐等阶段。上述开源能力经自主研发的中台层进行统一封装、编排与增强，"
        "形成可交付、可运营的企业级产品。"
    )
    doc.add_paragraph(body)

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.add_run("二、核心功能模块").bold = True

    modules = [
        (
            "1. 数据建设与治理（自主研发的数据工作台）",
            [
                "支持多格式文档（如 PDF、Markdown、DOCX、TXT 等）接入与版本管理；",
                "提供多种文本分块策略（结构感知、递归分隔、定长与代码感知等），保障下游训练语料粒度可控；",
                "面向领域构建标签体系与合成指令数据，支持单轮/多轮对话数据生产；",
                "内置人机协同审核与修订流程，保证数据集可追溯、可质检、可复现。",
            ],
        ),
        (
            "2. 预训练与继续预训练（CPT）",
            [
                "基于平台治理后的大规模语料，配置继续预训练或领域预训练任务；",
                "与训练引擎对接，支持从通用基座向行业分布迁移，提升领域语言建模能力；",
                "支持任务队列、资源与实验记录，便于多项目并行与结果对比。",
            ],
        ),
        (
            "3. 监督微调（SFT）",
            [
                "将平台产出的高质量指令数据与对话数据，一键对接微调流水线；",
                "支持全参数微调与高效参数微调等常见策略（以实际部署的训练框架能力为准）；",
                "形成「数据版本—训练配置—模型 Checkpoint」的清晰对应关系，支撑迭代发布。",
            ],
        ),
        (
            "4. 强化学习与偏好对齐",
            [
                "基于平台构造的偏好对、排序数据或评测反馈，开展 DPO、RLHF/PPO 等对齐训练（以所选算法与引擎为准）；",
                "在监督微调模型之上进一步优化安全性、风格与业务规则符合度；",
                "与评测模块联动，支持对齐前后效果对比与回归测试。",
            ],
        ),
        (
            "5. 评测与闭环运营",
            [
                "支持构建客观题、主观题与开放域问答等评测集，并可组织人机盲测对比；",
                "可对接 Judge 模型或规则打分，实现自动化批量评测；",
                "贯穿预训练数据质量、各阶段模型效果，支撑「评测—发现问题—回流标注/再训练」的持续改进闭环。",
            ],
        ),
    ]

    for heading, bullets in modules:
        h = doc.add_paragraph()
        h.add_run(heading).bold = True
        for b in bullets:
            doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.add_run("三、平台价值").bold = True
    values = [
        "自主研发的一体化工作台，降低多工具拼接成本与数据泄露风险；",
        "数据资产与训练实验可追溯、可审计，满足企业内控与合规要求；",
        "四阶段能力（预训练、微调、强化学习、评测）贯通，缩短从业务资料到可用模型的周期。",
    ]
    for v in values:
        doc.add_paragraph(v, style="List Bullet")

    doc.add_paragraph()
    foot = doc.add_paragraph(
        "说明：本文档所述功能以实际部署版本与合同范围为准；涉及开源组件的部分以相应开源协议及版本说明为准。"
    )
    for r in foot.runs:
        r.font.size = Pt(9)
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    foot.paragraph_format.space_before = Pt(12)

    for para in doc.paragraphs:
        if not para.runs:
            continue
        for run in para.runs:
            if run.font.size is None or run.font.size == Pt(11):
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                if run.font.size is None:
                    run.font.size = Pt(12)

    out = r"d:\legal-redaction\qingjing-dataset-platform-intro.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
